"""
工具函数模块
包含文件处理、验证、转换等通用功能
"""

import os
import uuid
import aiofiles
from pathlib import Path
from typing import Optional, Tuple
from fastapi import UploadFile, HTTPException
from loguru import logger
from config import settings
import subprocess
from docx import Document
import requests
import edge_tts
import re
from pydub import AudioSegment
import jieba
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import re
from datetime import datetime

DEEPSEEK_API_URL = "https://api.deepseek.com/chat/completions"
DEEPSEEK_API_KEY = "sk-f42549d7d09049009c8bfbc74c341939"
DEEPSEEK_MODEL = "deepseek-chat"


def generate_unique_filename(prefix: str, original_filename: str) -> str:
    """
    生成唯一的文件名，格式为 prefix_UUID.扩展名
    Args:
        prefix: 文件名前缀
        original_filename: 原始文件名
    Returns:
        唯一的文件名
    """
    file_extension = Path(original_filename).suffix
    unique_id = str(uuid.uuid4())
    return f"{prefix}_{unique_id}{file_extension}"


def validate_file_extension(filename: str) -> bool:
    """
    验证文件扩展名是否允许
    
    Args:
        filename: 文件名
        
    Returns:
        是否允许的文件类型
    """
    file_extension = Path(filename).suffix.lower()
    return file_extension in settings.ALLOWED_EXTENSIONS


def validate_file_size(file_size: int) -> bool:
    """
    验证文件大小是否在允许范围内
    
    Args:
        file_size: 文件大小（字节）
        
    Returns:
        文件大小是否合法
    """
    return file_size <= settings.MAX_FILE_SIZE


async def save_uploaded_file(content: bytes, filename: str) -> str:
    """
    保存上传的文件到本地存储
    
    Args:
        content: 文件内容（二进制）
        filename: 要保存的文件名
        
    Returns:
        保存的文件路径
        
    Raises:
        HTTPException: 文件保存失败时抛出异常
    """
    try:
        file_path = os.path.join(settings.UPLOAD_DIR, filename)
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(content)
        logger.info(f"文件已保存: {file_path}")
        return file_path
    except Exception as e:
        logger.error(f"保存文件失败: {str(e)}")
        raise HTTPException(status_code=500, detail="文件保存失败")


async def validate_and_save_file(upload_file: UploadFile) -> Tuple[str, str, int]:
    """
    验证并保存上传的文件
    
    Args:
        upload_file: 上传的文件对象
        
    Returns:
        元组包含 (文件名, 文件路径, 文件大小)
        
    Raises:
        HTTPException: 验证或保存失败时抛出异常
    """
    # 检查文件名是否存在
    if not upload_file.filename:
        raise HTTPException(status_code=400, detail="文件名不能为空")
    # 验证文件扩展名
    if not validate_file_extension(upload_file.filename):
        raise HTTPException(
            status_code=400, 
            detail=f"不支持的文件类型。允许的类型: {', '.join(settings.ALLOWED_EXTENSIONS)}"
        )
    # 读取文件内容以获取大小
    content = await upload_file.read()
    file_size = len(content)
    # 验证文件大小
    if not validate_file_size(file_size):
        max_size_mb = settings.MAX_FILE_SIZE // (1024 * 1024)
        raise HTTPException(
            status_code=400,
            detail=f"文件大小超过限制。最大允许: {max_size_mb}MB"
        )
    # 生成唯一文件名
    prefix = os.path.splitext(upload_file.filename)[0] if upload_file.filename else "input"
    unique_filename = generate_unique_filename(prefix, upload_file.filename)
    # 保存文件
    file_path = await save_uploaded_file(content, unique_filename)
    return unique_filename, file_path, file_size


def get_file_type(filename: str) -> str:
    """
    根据文件名获取文件类型
    
    Args:
        filename: 文件名
        
    Returns:
        文件类型（扩展名，不包含点）
    """
    return Path(filename).suffix.lower().lstrip('.')


async def convert_docx_to_pdf(docx_file_path: str, output_filename: str) -> str:
    """
    使用 LibreOffice 将 DOCX 文件转换为 PDF（Linux 环境）
    Args:
        docx_file_path: DOCX 文件路径
        output_filename: 输出 PDF 文件名
    Returns:
        生成的 PDF 文件路径
    Raises:
        HTTPException: 转换失败时抛出异常
    """
    try:
        pdf_file_path = os.path.join(settings.OUTPUT_DIR, output_filename)
        # 使用libreoffice命令行转换
        cmd = [
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            "--outdir", settings.OUTPUT_DIR,
            docx_file_path
        ]
        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if result.returncode != 0:
            logger.error(f"LibreOffice转换失败: {result.stderr.decode()}")
            raise Exception(result.stderr.decode())
        # LibreOffice会自动用原文件名.pdf输出，重命名为output_filename
        base_pdf = os.path.splitext(os.path.basename(docx_file_path))[0] + ".pdf"
        base_pdf_path = os.path.join(settings.OUTPUT_DIR, base_pdf)
        if base_pdf_path != pdf_file_path:
            os.rename(base_pdf_path, pdf_file_path)
        logger.info(f"DOCX转PDF成功: {docx_file_path} -> {pdf_file_path}")
        return pdf_file_path
    except Exception as e:
        logger.error(f"DOCX转PDF失败: {str(e)}")
        raise HTTPException(status_code=500, detail="文件转换失败")


def get_file_size(file_path: str) -> int:
    """
    获取文件大小
    
    Args:
        file_path: 文件路径
        
    Returns:
        文件大小（字节）
    """
    try:
        return os.path.getsize(file_path)
    except OSError:
        return 0


def file_exists(file_path: str) -> bool:
    """
    检查文件是否存在
    
    Args:
        file_path: 文件路径
        
    Returns:
        文件是否存在
    """
    return os.path.exists(file_path)


def cleanup_file(file_path: str) -> bool:
    """
    清理文件
    
    Args:
        file_path: 要删除的文件路径
        
    Returns:
        删除是否成功
    """
    try:
        if file_exists(file_path):
            os.remove(file_path)
            logger.info(f"文件已删除: {file_path}")
            return True
        return False
    except Exception as e:
        logger.error(f"删除文件失败: {str(e)}")
        return False 


def read_docx_text(docx_path: str) -> str:
    """
    读取 docx 文件的全部文本内容，按段落拼接为字符串。
    Args:
        docx_path: docx 文件路径
    Returns:
        文件全部文本内容
    """
    doc = Document(docx_path)
    return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])


def write_docx_text(text: str, output_path: str):
    """
    将文本内容写入新的 docx 文件，每行为一个段落。
    Args:
        text: 要写入的文本内容
        output_path: 输出 docx 文件路径
    """
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    doc.save(output_path) 


def call_deepseek(prompt: str) -> str:
    """
    调用 DeepSeek 大模型 API，将 prompt 发送给模型并返回生成内容。
    Args:
        prompt: 输入的提示词
    Returns:
        生成的文本内容
    Raises:
        HTTPException: 调用失败时抛出
    """
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": DEEPSEEK_MODEL,
        "messages": [
            {"role": "system", "content": "你是播客文案生成助手。"},
            {"role": "user", "content": prompt}
        ],
        "stream": False
    }
    try:
        resp = requests.post(DEEPSEEK_API_URL, headers=headers, json=data, timeout=60)
        resp.raise_for_status()
        return resp.json()["choices"][0]["message"]["content"]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DeepSeek大模型生成失败: {str(e)}") 


async def text_to_speech_edge_tts(text: str, output_path: str, voice: str = "zh-CN-XiaoxiaoNeural"):
    """
    使用 edge-tts 将文本转换为音频文件（mp3）
    Args:
        text: 要转换的文本
        output_path: 输出音频文件路径
        voice: 语音名称，默认中文女声
    """
    communicate = edge_tts.Communicate(text, voice=voice)
    await communicate.save(output_path) 


def split_dialogue(text: str):
    """
    按照“**女生：**”“**男生:**”分段，返回[(角色, 内容), ...]，兼容中英文冒号。
    """
    pattern = r"\\*\\*([男女]生)[：:]\\*\\*"
    segments = []
    last_pos = 0
    last_speaker = None
    for match in re.finditer(pattern, text):
        if last_speaker is not None:
            content = text[last_pos:match.start()].strip()
            if content:
                segments.append((last_speaker, content))
        last_speaker = match.group(1)
        last_pos = match.end()
    if last_speaker is not None:
        content = text[last_pos:].strip()
        if content:
            segments.append((last_speaker, content))
    return segments


def concat_audios(audio_paths, output_path):
    """
    拼接多个 mp3 音频文件为一个
    """
    combined = AudioSegment.empty()
    for path in audio_paths:
        combined += AudioSegment.from_file(path)
    combined.export(output_path, format="mp3") 


def clean_markdown(text: str) -> str:
    """
    去除常见 Markdown 强调符号（如 **、*、__、`）
    """
    import re
    return re.sub(r'\*\*|\*|__|`', '', text) 


def split_daily_report_by_date(text: str):
    """
    按日期分割日报内容，返回 [(date, content), ...]
    支持 2024-05-01、2024年5月1日、05月01日、2024/05/01 等格式
    """
    # 先输出前200字符看看格式
    logger.info(f"文本前200字符: {text[:200]}")
    
    # 匹配多种日期格式，包括 YYYY.M.D 格式
    patterns = [
        r"(\d{4}\.\d{1,2}\.\d{1,2}(?:[（(][^）)]*[）)])?)",  # 2025.7.14 或 2025.7.14（一）
        r"(\d{4}年\d{1,2}月\d{1,2}日)",      # 2024年5月1日
        r"(\d{1,2}月\d{1,2}日)",              # 5月1日
        r"(\d{4}-\d{1,2}-\d{1,2})",          # 2024-05-01
        r"(\d{4}/\d{1,2}/\d{1,2})",          # 2024/05/01
        r"(\d{1,2}/\d{1,2})",                # 05/01
        r"(\d{1,2}-\d{1,2})",                # 05-01
    ]
    
    matches = []
    for i, pattern in enumerate(patterns):
        pattern_matches = list(re.finditer(pattern, text))
        if pattern_matches:
            matches.extend(pattern_matches)
            logger.info(f"模式 {i+1} 找到 {len(pattern_matches)} 个匹配: {[m.group(1) for m in pattern_matches[:3]]}")
    
    # 按位置排序
    matches.sort(key=lambda x: x.start())
    logger.info(f"总共找到 {len(matches)} 个日期匹配: {[m.group(1) for m in matches[:5]]}")
    
    results = []
    for i, match in enumerate(matches):
        date_str = match.group(1)
        start = match.end()
        end = matches[i+1].start() if i+1 < len(matches) else len(text)
        content = text[start:end].strip()
        # 规范化日期
        date = normalize_date(date_str)
        if content:
            results.append((date, content))
            logger.info(f"添加日期段落: {date} (内容长度: {len(content)})")
    return results

def normalize_date(date_str: str) -> str:
    """
    尝试多种格式转为 yyyy-mm-dd
    """
    # 移除中文括号和星期
    date_str = re.sub(r'（[一二三四五六日]）', '', date_str)
    
    # 处理 YYYY.M.D 格式
    if re.match(r"^\\d{4}\\.\\d{1,2}\\.\\d{1,2}$", date_str):
        date_str = date_str.replace('.', '-')
    
    date_str = date_str.replace('年', '-').replace('月', '-').replace('日', '').replace('号', '')
    date_str = date_str.replace('/', '-')
    
    # 处理无年份情况，补今年
    if re.match(r"^\\d{1,2}-\\d{1,2}$", date_str):
        date_str = f"{datetime.now().year}-{date_str}"
    
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        return dt.strftime("%Y-%m-%d")
    except:
        logger.warning(f"无法解析日期格式: {date_str}")
        return date_str  # 原样返回

def summarize_content_with_deepseek(date, content):
    """
    调用deepseek播客文稿模块，对单日内容生成一句话摘要
    """
    prompt = f"""
你是一名专业AI新闻播客编剧，请用一句话总结{date}这一天AI产品日报的主要内容，突出重点和亮点，简明扼要，风格口语化但不失专业。原文如下：\n{content}
"""
    return call_deepseek(prompt).strip()

def generate_wordcloud(text: str, output_path: str):
    """
    生成中文词云图片，使用哈工大停用词库，只显示名词
    """
    # 文本预处理：过滤无效字符
    import re
    # 移除数字、特殊符号、URL等
    text = re.sub(r'\d+', '', text)  # 移除数字
    text = re.sub(r'https?://\S+', '', text)  # 移除URL
    text = re.sub(r'[^\u4e00-\u9fa5a-zA-Z\s]', '', text)  # 只保留中文、英文、空格
    
    # 加载哈工大停用词库
    stopwords = set()
    try:
        with open('hit_stopwords.txt', 'r', encoding='utf-8') as f:
            stopwords = set(line.strip() for line in f)
    except FileNotFoundError:
        # 如果文件不存在，使用基本停用词
        stopwords = {'的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个', '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看', '好', '自己', '这'}
    
    # 使用jieba进行分词和词性标注，只保留名词
    import jieba.posseg as pseg
    words = []
    for word, flag in pseg.cut(text):
        # 只保留名词 (n开头) 且不在停用词中的词
        if flag.startswith('n') and word not in stopwords and len(word) > 1:
            words.append(word)
    
    # 如果没有有效词汇，使用所有词汇
    if not words:
        words = [word for word in jieba.cut(text) if word not in stopwords and len(word) > 1]
    
    words_text = " ".join(words)
    
    # 使用中文字体
    wc = WordCloud(
        font_path="/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        width=800, 
        height=400, 
        background_color="white"
    )
    wc.generate(words_text)
    wc.to_file(output_path) 
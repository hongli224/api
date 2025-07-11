"""
API测试脚本
用于测试File Conversion API的各个功能
"""

import requests
import json
import time
from pathlib import Path


class APITester:
    """API测试类"""
    
    def __init__(self, base_url="http://localhost:5000"):
        """
        初始化测试器
        
        Args:
            base_url: API服务的基础URL
        """
        self.base_url = base_url
        self.session = requests.Session()
    
    def test_health_check(self):
        """测试健康检查接口"""
        print("🔍 测试健康检查接口...")
        try:
            response = self.session.get(f"{self.base_url}/health")
            print(f"状态码: {response.status_code}")
            print(f"响应: {response.json()}")
            return response.status_code == 200
        except Exception as e:
            print(f"❌ 健康检查失败: {str(e)}")
            return False
    
    def test_root_endpoint(self):
        """测试根路径接口"""
        print("\n🔍 测试根路径接口...")
        try:
            response = self.session.get(f"{self.base_url}/")
            print(f"状态码: {response.status_code}")
            print(f"响应: {response.json()}")
            return response.status_code == 200
        except Exception as e:
            print(f"❌ 根路径测试失败: {str(e)}")
            return False
    
    def test_file_upload(self, file_path):
        """
        测试文件上传接口
        
        Args:
            file_path: 要上传的文件路径
            
        Returns:
            上传的文件ID或None
        """
        print(f"\n📁 测试文件上传接口: {file_path}")
        
        if not Path(file_path).exists():
            print(f"❌ 文件不存在: {file_path}")
            return None
        
        try:
            with open(file_path, 'rb') as f:
                files = {'file': (Path(file_path).name, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
                response = self.session.post(f"{self.base_url}/api/file_upload", files=files)
            
            print(f"状态码: {response.status_code}")
            if response.status_code == 201:
                result = response.json()
                print(f"✅ 文件上传成功: {result}")
                return result.get('id')
            else:
                print(f"❌ 文件上传失败: {response.text}")
                return None
                
        except Exception as e:
            print(f"❌ 文件上传异常: {str(e)}")
            return None
    
    def test_get_files(self):
        """测试获取文件列表接口"""
        print("\n📋 测试获取文件列表接口...")
        try:
            response = self.session.get(f"{self.base_url}/api/files")
            print(f"状态码: {response.status_code}")
            if response.status_code == 200:
                files = response.json()
                print(f"✅ 获取到 {len(files)} 个文件")
                for file in files:
                    print(f"  - {file['original_filename']} (ID: {file['id']})")
                return files
            else:
                print(f"❌ 获取文件列表失败: {response.text}")
                return []
        except Exception as e:
            print(f"❌ 获取文件列表异常: {str(e)}")
            return []
    
    def test_get_file(self, file_id):
        """
        测试获取单个文件接口
        
        Args:
            file_id: 文件ID
        """
        print(f"\n📄 测试获取单个文件接口: {file_id}")
        try:
            response = self.session.get(f"{self.base_url}/api/files/{file_id}")
            print(f"状态码: {response.status_code}")
            if response.status_code == 200:
                file_info = response.json()
                print(f"✅ 文件信息: {file_info}")
                return file_info
            else:
                print(f"❌ 获取文件信息失败: {response.text}")
                return None
        except Exception as e:
            print(f"❌ 获取文件信息异常: {str(e)}")
            return None
    
    def test_convert_docx_to_pdf(self, file_id):
        """
        测试DOCX转PDF接口
        
        Args:
            file_id: 要转换的文件ID
        """
        print(f"\n🔄 测试DOCX转PDF接口: {file_id}")
        try:
            data = {"file_id": file_id}
            response = self.session.post(
                f"{self.base_url}/api/convert_docx_to_pdf",
                json=data,
                headers={'Content-Type': 'application/json'}
            )
            print(f"状态码: {response.status_code}")
            if response.status_code == 201:
                result = response.json()
                print(f"✅ 转换成功: {result}")
                return result
            else:
                print(f"❌ 转换失败: {response.text}")
                return None
        except Exception as e:
            print(f"❌ 转换异常: {str(e)}")
            return None
    
    def test_delete_file(self, file_id):
        """
        测试删除文件接口
        
        Args:
            file_id: 要删除的文件ID
        """
        print(f"\n🗑️ 测试删除文件接口: {file_id}")
        try:
            response = self.session.delete(f"{self.base_url}/api/files/{file_id}")
            print(f"状态码: {response.status_code}")
            if response.status_code == 200:
                result = response.json()
                print(f"✅ 删除成功: {result}")
                return True
            else:
                print(f"❌ 删除失败: {response.text}")
                return False
        except Exception as e:
            print(f"❌ 删除异常: {str(e)}")
            return False
    
    def run_full_test(self, test_file_path=None):
        """
        运行完整的API测试
        
        Args:
            test_file_path: 测试文件路径（可选）
        """
        print("🚀 开始运行API测试套件...")
        print("=" * 50)
        
        # 测试基础接口
        self.test_health_check()
        self.test_root_endpoint()
        
        # 测试文件管理接口
        files = self.test_get_files()
        
        # 如果有测试文件，测试上传和转换
        if test_file_path:
            file_id = self.test_file_upload(test_file_path)
            if file_id:
                self.test_get_file(file_id)
                self.test_convert_docx_to_pdf(file_id)
                
                # 等待转换完成
                print("\n⏳ 等待转换完成...")
                time.sleep(2)
                
                # 再次获取文件列表查看转换结果
                self.test_get_files()
        
        print("\n" + "=" * 50)
        print("✅ API测试完成！")


def create_test_docx():
    """
    创建一个简单的测试DOCX文件
    注意：这需要python-docx库
    """
    try:
        from docx import Document
        
        # 创建文档
        doc = Document()
        doc.add_heading('测试文档', 0)
        doc.add_paragraph('这是一个用于测试API的DOCX文档。')
        doc.add_paragraph('包含一些基本的文本内容。')
        
        # 保存文档
        test_file_path = "test_document.docx"
        doc.save(test_file_path)
        print(f"✅ 测试文档已创建: {test_file_path}")
        return test_file_path
        
    except ImportError:
        print("❌ python-docx库未安装，无法创建测试文档")
        return None
    except Exception as e:
        print(f"❌ 创建测试文档失败: {str(e)}")
        return None


if __name__ == "__main__":
    """主测试函数"""
    print("🧪 File Conversion API 测试工具")
    print("=" * 50)
    
    # 创建测试器
    tester = APITester()
    
    # 检查服务是否运行
    if not tester.test_health_check():
        print("❌ API服务未运行，请先启动服务")
        exit(1)
    
    # 创建测试文档
    test_file = create_test_docx()
    
    # 运行完整测试
    tester.run_full_test(test_file)
    
    # 清理测试文件
    if test_file and Path(test_file).exists():
        Path(test_file).unlink()
        print(f"🧹 已清理测试文件: {test_file}") 